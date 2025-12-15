"""
Word Document Generator Agent
Generates a professional Word document for infrastructure solution architecture
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import json
from datetime import datetime


class WordDocumentAgent:
    """Generates professional Word documents from POC outputs"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Configure document styles"""
        # Set normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Set heading styles
        for i in range(1, 4):
            style = self.doc.styles[f'Heading {i}']
            font = style.font
            font.name = 'Calibri'
            font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    def generate(self, rfp_file: str, arch_file: str, cost_file: str, 
                 proposal_file: str, roadmap_file: str, diagram_image: str = None):
        """
        Generate comprehensive Word document
        
        Args:
            rfp_file: Path to rfp_analysis.json
            arch_file: Path to architecture.json
            cost_file: Path to cost.json
            proposal_file: Path to proposal.md
            roadmap_file: Path to roadmap.json
            diagram_image: Optional path to diagram PNG
        """
        # Load data
        rfp = self._load_json(rfp_file)
        arch = self._load_json(arch_file)
        cost = self._load_json(cost_file)
        roadmap = self._load_json(roadmap_file)
        
        # Build document sections
        self._add_title_page()
        self._add_executive_summary(rfp, cost)
        self._add_table_of_contents()
        self._add_business_requirements(rfp)
        self._add_proposed_architecture(arch, diagram_image)
        self._add_technical_specifications(arch)
        self._add_security_compliance(arch)
        self._add_cost_breakdown(cost)
        self._add_implementation_roadmap(roadmap)
        self._add_assumptions_risks()
        self._add_appendix()
        
        return self.doc
    
    def _load_json(self, filepath: str) -> dict:
        """Load JSON file"""
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f)
    
    def _add_title_page(self):
        """Add professional title page"""
        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('INFRASTRUCTURE SOLUTION ARCHITECTURE')
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Cloud Migration & Modernization Proposal')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        self.doc.add_paragraph()  # Spacing
        
        # Client info
        client = self.doc.add_paragraph()
        client.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = client.add_run(f'Prepared for: MMI Cloud Migration Project')
        run.font.size = Pt(14)
        
        # Date
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}')
        run.font.size = Pt(12)
        
        # Page break
        self.doc.add_page_break()
    
    def _add_executive_summary(self, rfp: dict, cost: dict):
        """Add executive summary section"""
        self.doc.add_heading('Executive Summary', 1)
        
        summary = self.doc.add_paragraph()
        summary.add_run(
            f"This document presents a comprehensive infrastructure solution architecture "
            f"for the cloud migration and modernization initiative. "
        )
        
        self.doc.add_heading('Business Objectives', 2)
        for goal in rfp.get('business_goals', [])[:5]:
            self.doc.add_paragraph(goal, style='List Bullet')
        
        self.doc.add_heading('Investment Summary', 2)
        total_monthly = cost.get('summary', {}).get('total_monthly', 0)
        total_annual = cost.get('summary', {}).get('total_annual', 0)
        
        table = self.doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.rows[0].cells[0].text = 'Estimated Monthly Cost'
        table.rows[0].cells[1].text = f'${total_monthly:,.2f}'
        table.rows[1].cells[0].text = 'Estimated Annual Cost'
        table.rows[1].cells[1].text = f'${total_annual:,.2f}'
        table.rows[2].cells[0].text = 'Implementation Timeline'
        table.rows[2].cells[1].text = '6-9 months'
        
        self.doc.add_page_break()
    
    def _add_table_of_contents(self):
        """Add table of contents placeholder"""
        self.doc.add_heading('Table of Contents', 1)
        self.doc.add_paragraph(
            'Note: In Microsoft Word, right-click here and select '
            '"Update Field" to generate the table of contents.'
        )
        self.doc.add_page_break()
    
    def _add_business_requirements(self, rfp: dict):
        """Add business requirements section"""
        self.doc.add_heading('1. Business Requirements Analysis', 1)
        
        # Business Goals
        self.doc.add_heading('1.1 Business Goals', 2)
        for goal in rfp.get('business_goals', []):
            self.doc.add_paragraph(goal, style='List Bullet')
        
        # Functional Requirements
        self.doc.add_heading('1.2 Functional Requirements', 2)
        for req in rfp.get('functional_requirements', [])[:10]:
            self.doc.add_paragraph(req, style='List Bullet')
        
        # Non-Functional Requirements
        self.doc.add_heading('1.3 Non-Functional Requirements', 2)
        nfr = rfp.get('non_functional_requirements', {})
        
        if nfr.get('performance'):
            self.doc.add_paragraph(f"Performance: {nfr['performance']}", style='List Bullet')
        if nfr.get('security'):
            self.doc.add_paragraph(f"Security: {nfr['security']}", style='List Bullet')
        if nfr.get('availability'):
            self.doc.add_paragraph(f"Availability: {nfr['availability']}", style='List Bullet')
        
        self.doc.add_page_break()
    
    def _add_proposed_architecture(self, arch: dict, diagram_image: str = None):
        """Add proposed architecture section"""
        self.doc.add_heading('2. Proposed Solution Architecture', 1)
        
        self.doc.add_heading('2.1 Architecture Overview', 2)
        self.doc.add_paragraph(
            f"The proposed solution follows a modern cloud-native architecture "
            f"with {len(arch.get('layers', []))} distinct layers designed for "
            f"scalability, security, and operational excellence."
        )
        
        # Add diagram if available
        if diagram_image and Path(diagram_image).exists():
            self.doc.add_heading('2.2 Architecture Diagram', 2)
            try:
                self.doc.add_picture(diagram_image, width=Inches(6))
                last_paragraph = self.doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                self.doc.add_paragraph(f'Diagram file not found: {diagram_image}')
        
        # Architecture Layers
        self.doc.add_heading('2.3 Architecture Layers', 2)
        for layer in arch.get('layers', []):
            self.doc.add_heading(f"{layer['name']}", 3)
            
            # Components
            self.doc.add_paragraph('Components:', style='List Bullet')
            for component in layer.get('components', []):
                p = self.doc.add_paragraph(component, style='List Bullet 2')
            
            # Responsibilities
            self.doc.add_paragraph('Responsibilities:', style='List Bullet')
            for resp in layer.get('responsibilities', []):
                self.doc.add_paragraph(resp, style='List Bullet 2')
        
        self.doc.add_page_break()
    
    def _add_technical_specifications(self, arch: dict):
        """Add technical specifications"""
        self.doc.add_heading('3. Technical Specifications', 1)
        
        # Data Flow
        self.doc.add_heading('3.1 Data Flow Architecture', 2)
        flows = arch.get('data_flow', [])
        if flows:
            table = self.doc.add_table(rows=len(flows)+1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Header
            table.rows[0].cells[0].text = 'Source'
            table.rows[0].cells[1].text = 'Destination'
            table.rows[0].cells[2].text = 'Protocol'
            
            # Data
            for i, flow in enumerate(flows, 1):
                table.rows[i].cells[0].text = flow.get('from', '')
                table.rows[i].cells[1].text = flow.get('to', '')
                table.rows[i].cells[2].text = flow.get('protocol', '')
        
        # Scalability
        self.doc.add_heading('3.2 Scalability Approach', 2)
        self.doc.add_paragraph(arch.get('scalability_approach', 'Not specified'))
        
        # Disaster Recovery
        if 'disaster_recovery' in arch:
            self.doc.add_heading('3.3 Disaster Recovery', 2)
            dr = arch['disaster_recovery']
            self.doc.add_paragraph(f"Recovery Point Objective (RPO): {dr.get('rpo', 'TBD')}")
            self.doc.add_paragraph(f"Recovery Time Objective (RTO): {dr.get('rto', 'TBD')}")
            if dr.get('strategy'):
                self.doc.add_paragraph(f"Strategy: {dr['strategy']}")
        
        self.doc.add_page_break()
    
    def _add_security_compliance(self, arch: dict):
        """Add security and compliance section"""
        self.doc.add_heading('4. Security & Compliance', 1)
        
        self.doc.add_heading('4.1 Security Controls', 2)
        for control in arch.get('security_controls', []):
            self.doc.add_paragraph(control, style='List Bullet')
        
        self.doc.add_heading('4.2 Compliance Requirements', 2)
        self.doc.add_paragraph(
            'The proposed architecture is designed to meet the following compliance standards:'
        )
        self.doc.add_paragraph('Enterprise EDS (Encryption Data Standards)', style='List Bullet')
        self.doc.add_paragraph('Data Protection and Privacy Regulations', style='List Bullet')
        self.doc.add_paragraph('Industry Best Practices for Cloud Security', style='List Bullet')
        
        self.doc.add_page_break()
    
    def _add_cost_breakdown(self, cost: dict):
        """Add cost breakdown section"""
        self.doc.add_heading('5. Cost Analysis', 1)
        
        self.doc.add_heading('5.1 Detailed Cost Breakdown', 2)
        
        line_items = cost.get('line_items', [])
        if line_items:
            table = self.doc.add_table(rows=len(line_items)+2, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header
            header = table.rows[0].cells
            header[0].text = 'Service'
            header[1].text = 'Category'
            header[2].text = 'Monthly Cost'
            header[3].text = 'Annual Cost'
            
            # Data
            for i, item in enumerate(line_items, 1):
                row = table.rows[i].cells
                row[0].text = item.get('service', '')
                row[1].text = item.get('category', '')
                monthly = item.get('monthly', 0)
                row[2].text = f'${monthly:,.2f}'
                row[3].text = f'${monthly * 12:,.2f}'
            
            # Total
            summary = cost.get('summary', {})
            total_row = table.rows[-1].cells
            total_row[0].text = 'TOTAL'
            total_row[0].merge(total_row[1])
            total_row[2].text = f"${summary.get('total_monthly', 0):,.2f}"
            total_row[3].text = f"${summary.get('total_annual', 0):,.2f}"
            
            # Bold the total row
            for cell in table.rows[-1].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        self.doc.add_heading('5.2 Cost Optimization Opportunities', 2)
        self.doc.add_paragraph('Reserved Instances for predictable workloads (up to 40% savings)', style='List Bullet')
        self.doc.add_paragraph('Auto-scaling to match demand and reduce waste', style='List Bullet')
        self.doc.add_paragraph('S3 lifecycle policies for archival data storage', style='List Bullet')
        self.doc.add_paragraph('Right-sizing based on actual usage patterns', style='List Bullet')
        
        self.doc.add_page_break()
    
    def _add_implementation_roadmap(self, roadmap: dict):
        """Add implementation roadmap section"""
        self.doc.add_heading('6. Implementation Roadmap', 1)
        
        phases = roadmap.get('phases', [])
        for phase in phases:
            self.doc.add_heading(f"{phase.get('name', 'Phase')}", 2)
            self.doc.add_paragraph(f"Duration: {phase.get('duration', 'TBD')}")
            
            if phase.get('deliverables'):
                self.doc.add_paragraph('Deliverables:', style='List Bullet')
                for deliverable in phase['deliverables']:
                    self.doc.add_paragraph(deliverable, style='List Bullet 2')
            
            if phase.get('milestones'):
                self.doc.add_paragraph('Key Milestones:', style='List Bullet')
                for milestone in phase['milestones']:
                    self.doc.add_paragraph(milestone, style='List Bullet 2')
        
        self.doc.add_page_break()
    
    def _add_assumptions_risks(self):
        """Add assumptions and risks section"""
        self.doc.add_heading('7. Assumptions & Risks', 1)
        
        self.doc.add_heading('7.1 Key Assumptions', 2)
        self.doc.add_paragraph('Existing infrastructure will remain operational during migration', style='List Bullet')
        self.doc.add_paragraph('Necessary access and permissions will be provided timely', style='List Bullet')
        self.doc.add_paragraph('SME resources will be available for knowledge transfer', style='List Bullet')
        self.doc.add_paragraph('Network bandwidth is sufficient for data migration', style='List Bullet')
        
        self.doc.add_heading('7.2 Risk Mitigation', 2)
        
        risks_table = self.doc.add_table(rows=4, cols=3)
        risks_table.style = 'Light Grid Accent 1'
        
        # Header
        risks_table.rows[0].cells[0].text = 'Risk'
        risks_table.rows[0].cells[1].text = 'Impact'
        risks_table.rows[0].cells[2].text = 'Mitigation Strategy'
        
        # Risks
        risks_table.rows[1].cells[0].text = 'Data Migration Delays'
        risks_table.rows[1].cells[1].text = 'High'
        risks_table.rows[1].cells[2].text = 'Parallel runs, incremental migration approach'
        
        risks_table.rows[2].cells[0].text = 'Performance Issues'
        risks_table.rows[2].cells[1].text = 'Medium'
        risks_table.rows[2].cells[2].text = 'Load testing, auto-scaling, performance monitoring'
        
        risks_table.rows[3].cells[0].text = 'Cost Overruns'
        risks_table.rows[3].cells[1].text = 'Medium'
        risks_table.rows[3].cells[2].text = 'Cost monitoring, budget alerts, reserved instances'
        
        self.doc.add_page_break()
    
    def _add_appendix(self):
        """Add appendix section"""
        self.doc.add_heading('8. Appendix', 1)
        
        self.doc.add_heading('8.1 Glossary', 2)
        terms = [
            ('AWS', 'Amazon Web Services'),
            ('DMS', 'Database Migration Service'),
            ('EMR', 'Elastic MapReduce'),
            ('ECS', 'Elastic Container Service'),
            ('RDS', 'Relational Database Service'),
            ('S3', 'Simple Storage Service'),
            ('KMS', 'Key Management Service'),
            ('VPC', 'Virtual Private Cloud'),
        ]
        
        table = self.doc.add_table(rows=len(terms)+1, cols=2)
        table.style = 'Light Grid Accent 1'
        table.rows[0].cells[0].text = 'Term'
        table.rows[0].cells[1].text = 'Definition'
        
        for i, (term, definition) in enumerate(terms, 1):
            table.rows[i].cells[0].text = term
            table.rows[i].cells[1].text = definition
        
        self.doc.add_heading('8.2 References', 2)
        self.doc.add_paragraph('AWS Well-Architected Framework', style='List Bullet')
        self.doc.add_paragraph('AWS Cloud Adoption Framework', style='List Bullet')
        self.doc.add_paragraph('Enterprise Architecture Standards', style='List Bullet')
    
    def save(self, output_path: str):
        """Save document to file"""
        self.doc.save(output_path)
        print(f"✓ Word document saved: {output_path}")


def generate_word_document(output_file: str = "output/Infrastructure_Solution_Architecture.docx"):
    """
    Generate Word document from POC outputs
    
    Args:
        output_file: Path to save the Word document
    """
    try:
        agent = WordDocumentAgent()
        
        # Generate document from output files
        agent.generate(
            rfp_file="output/rfp_analysis.json",
            arch_file="output/architecture.json",
            cost_file="output/cost.json",
            proposal_file="output/proposal.md",
            roadmap_file="output/roadmap.json",
            diagram_image="mmi_complete_architecture.png"  # Optional
        )
        
        # Save
        agent.save(output_file)
        
        return output_file
        
    except Exception as e:
        print(f"⚠️ Error generating Word document: {e}")
        return None


if __name__ == "__main__":
    # Test standalone
    generate_word_document()
